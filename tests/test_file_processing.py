import sys
import os
import io

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "app"))

import file_processing as fp


class TestExtractText:
    def test_extracts_plain_text(self):
        raw = "Hello, this is a plain text file.\nSecond line.".encode("utf-8")
        result = fp.extract_text("test.txt", "text/plain", raw)
        assert "Hello, this is a plain text file." in result
        assert "Second line" in result

    def test_extracts_docx(self):
        import docx

        doc = docx.Document()
        doc.add_paragraph("First paragraph of the document.")
        doc.add_paragraph("Second paragraph with more content.")
        buf = io.BytesIO()
        doc.save(buf)

        result = fp.extract_text(
            "test.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            buf.getvalue(),
        )
        assert "First paragraph of the document" in result
        assert "Second paragraph" in result

    def test_extracts_pdf(self):
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 750, "This is a test PDF.")
        c.drawString(100, 730, "Second line of PDF text.")
        c.save()

        result = fp.extract_text("test.pdf", "application/pdf", buf.getvalue())
        assert "This is a test PDF" in result
        assert "Second line" in result

    def test_rejects_unsupported_content_type(self):
        try:
            fp.extract_text("test.exe", "application/x-msdownload", b"fake bytes")
            assert False, "should have raised UnsupportedFileTypeError"
        except fp.UnsupportedFileTypeError as exc:
            assert "unsupported content type" in str(exc)

    def test_truncates_very_long_text(self):
        long_text = ("word " * 20_000).encode("utf-8")  # ~100,000 chars
        result = fp.extract_text("big.txt", "text/plain", long_text)
        assert len(result) <= fp.MAX_EXTRACTED_CHARS

    def test_corrupt_pdf_raises_extraction_error(self):
        garbage_bytes = b"this is not a real pdf file at all, just garbage bytes %PDF-fake"
        try:
            fp.extract_text("corrupt.pdf", "application/pdf", garbage_bytes)
            assert False, "should have raised ExtractionError"
        except fp.ExtractionError as exc:
            assert "failed to extract text" in str(exc)